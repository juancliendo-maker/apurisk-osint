"""Reporte Ejecutivo Diario — APURISK 1.0

Diseño visual de 2-3 páginas con foco en ANÁLISIS DE TENDENCIAS:
  Página 1: Portada · KPIs · Síntesis ejecutiva · Score gauge
  Página 2: Top 5 factores P×I · Casos persistentes (tendencias 7d)
  Página 3: Top alertas críticas con acciones · URLs clave

Genera dos formatos:
  - DOCX (python-docx) profesional, color-coded
  - PDF (reportlab) con misma estructura visual
"""
from __future__ import annotations
import json
import glob
from datetime import datetime, timedelta
from pathlib import Path
from collections import defaultdict

# -------------------- Helpers comunes --------------------
_MESES = {1: "ene", 2: "feb", 3: "mar", 4: "abr", 5: "may", 6: "jun",
          7: "jul", 8: "ago", 9: "set", 10: "oct", 11: "nov", 12: "dic"}


def _fmt_dt(iso_str: str | None) -> str:
    if not iso_str:
        return "—"
    try:
        s = iso_str.replace("Z", "+00:00")
        dt = datetime.fromisoformat(s)
        return f"{dt.day:02d} {_MESES.get(dt.month, '?')} {dt.strftime('%H:%M')} PET"
    except Exception:
        return iso_str[:16]


def _fmt_dt_full(iso_str: str | None) -> str:
    if not iso_str:
        return "—"
    try:
        s = iso_str.replace("Z", "+00:00")
        dt = datetime.fromisoformat(s)
        return f"{dt.day:02d} de {_MESES.get(dt.month, '?')} de {dt.year} · {dt.strftime('%H:%M')} PET"
    except Exception:
        return iso_str


def _aggregate_tendencias(snapshot_dir: str) -> dict:
    """Calcula tendencias de 7 días leyendo todos los snapshots disponibles."""
    p = Path(snapshot_dir)
    cutoff = datetime.now().astimezone() - timedelta(days=7)
    snaps = []
    for f in sorted(p.glob("apurisk_snapshot_*.json")):
        try:
            with open(f, encoding="utf-8") as fh:
                data = json.load(fh)
            try:
                gen = data.get("generado", "")
                dt = datetime.fromisoformat(gen.replace("Z", "+00:00"))
                if dt.tzinfo is None:
                    from datetime import timezone, timedelta as td
                    dt = dt.replace(tzinfo=timezone(td(hours=-5)))
                if dt >= cutoff:
                    snaps.append((dt, data))
            except Exception:
                continue
        except Exception:
            continue

    if not snaps:
        return {"snapshots": 0, "score_avg": 0, "score_actual": 0,
                "delta_vs_avg": 0, "alertas_persistentes": [],
                "factores_subiendo": [], "factores_bajando": []}

    last = snaps[-1][1]
    score_actual = last.get("riesgo", {}).get("global", 0)
    scores = [s[1].get("riesgo", {}).get("global", 0) for s in snaps]
    score_avg = sum(scores) / len(scores) if scores else 0
    score_max = max(scores) if scores else 0
    score_min = min(scores) if scores else 0

    # Casos persistentes
    grupos = defaultdict(list)
    for dt, data in snaps:
        for a in data.get("alertas", []):
            key = a.get("titulo", "")[:90].lower().strip()
            grupos[key].append((dt, a))
    persistentes = []
    for key, ocurrencias in grupos.items():
        if len(ocurrencias) < 2:
            continue
        last_a = ocurrencias[-1][1]
        timestamps = [o[0] for o in ocurrencias]
        dias = len({t.date() for t in timestamps})
        persistentes.append({
            "titulo": last_a.get("titulo", ""),
            "categoria": last_a.get("categoria", ""),
            "nivel": last_a.get("nivel", ""),
            "url": last_a.get("url", ""),
            "ocurrencias": len(ocurrencias),
            "dias_activo": dias,
            "primera_vez": min(timestamps).isoformat(timespec="seconds"),
        })
    persistentes.sort(key=lambda x: (-x["dias_activo"], -x["ocurrencias"]))

    # Factores que suben/bajan
    factores_serie = defaultdict(list)
    for dt, data in snaps:
        for f in data.get("matriz_riesgo", []):
            factores_serie[f["nombre"]].append(f["score"])

    subiendo, bajando = [], []
    for nombre, serie in factores_serie.items():
        if len(serie) < 2:
            continue
        delta = serie[-1] - serie[0]
        if delta >= 5:
            subiendo.append({"nombre": nombre, "delta": round(delta, 1),
                             "actual": round(serie[-1], 1)})
        elif delta <= -5:
            bajando.append({"nombre": nombre, "delta": round(delta, 1),
                            "actual": round(serie[-1], 1)})
    subiendo.sort(key=lambda x: -x["delta"])
    bajando.sort(key=lambda x: x["delta"])

    return {
        "snapshots": len(snaps),
        "score_avg": round(score_avg, 1),
        "score_actual": round(score_actual, 1),
        "score_max": round(score_max, 1),
        "score_min": round(score_min, 1),
        "delta_vs_avg": round(score_actual - score_avg, 1),
        "alertas_persistentes": persistentes[:6],
        "factores_subiendo": subiendo[:5],
        "factores_bajando": bajando[:3],
    }


# ============================================================================
# WORD (DOCX)
# ============================================================================
def generar_ejecutivo_docx(output_path: str, snapshot: dict, snapshot_dir: str) -> str:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    AZUL = RGBColor(0x1E, 0x40, 0xAF)
    AZUL_OSCURO = RGBColor(0x0F, 0x17, 0x2A)
    GRIS = RGBColor(0x64, 0x74, 0x8B)
    ROJO = RGBColor(0xDC, 0x26, 0x26)
    NARANJA = RGBColor(0xEA, 0x58, 0x0C)
    AMARILLO = RGBColor(0xD4, 0x8C, 0x0B)
    VERDE = RGBColor(0x16, 0xA3, 0x4A)

    def color_nivel(nivel):
        return {"CRÍTICO": ROJO, "CRÍTICA": ROJO, "ALTO": NARANJA, "ALTA": NARANJA,
                "MEDIO": AMARILLO, "MEDIA": AMARILLO, "BAJO": VERDE, "BAJA": VERDE}.get(nivel, GRIS)

    def shade_cell(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def borderless(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "nil")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    doc = Document()
    # Márgenes ajustados para mejor uso del espacio
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    riesgo = snapshot.get("riesgo", {})
    score = riesgo.get("global", 0)
    nivel = riesgo.get("nivel", "—")
    alertas = snapshot.get("alertas", [])
    crit_count = len([a for a in alertas if a["nivel"] == "CRÍTICA"])
    matriz = snapshot.get("matriz_riesgo", [])
    fecha_gen = snapshot.get("generado", "")
    art_24h = snapshot.get("n_articulos_24h", 0)
    n_tweets = snapshot.get("n_tweets", 0)
    n_conf = snapshot.get("n_conflictos", 0)
    sentimiento = riesgo.get("sentimiento_promedio", 0)

    # Tendencias
    tend = _aggregate_tendencias(snapshot_dir)

    # ====== PORTADA / HEADER ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("APURISK 1.0")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = AZUL_OSCURO
    p.add_run("  ").font.size = Pt(22)
    r2 = p.add_run("· Reporte Ejecutivo Diario")
    r2.italic = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = AZUL

    p = doc.add_paragraph()
    r = p.add_run(f"Riesgo Político · Perú  ·  {_fmt_dt_full(fecha_gen)}")
    r.font.size = Pt(10)
    r.font.color.rgb = GRIS

    # Línea divisoria visual (border bottom de un párrafo vacío)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "1E40AF")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ====== KPIs en tabla 4x1 ======
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kpi_data = [
        ("SCORE GLOBAL", f"{score}", nivel, color_nivel(nivel)),
        ("ALERTAS CRÍTICAS", f"{crit_count}", f"de {len(alertas)} totales", ROJO),
        ("CONFLICTOS SOC.", f"{n_conf}", "activos", NARANJA),
        ("COBERTURA 24H", f"{art_24h + n_tweets}", f"{art_24h} medios + {n_tweets} X", AZUL),
    ]
    for i, (lbl, val, sub, col) in enumerate(kpi_data):
        cell = kpi_table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "F8FAFC")
        # Label arriba
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(lbl)
        r.font.size = Pt(8)
        r.font.color.rgb = GRIS
        r.bold = True
        # Valor grande
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(val)
        r.font.size = Pt(28)
        r.font.color.rgb = col
        r.bold = True
        # Sub
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p3.add_run(sub)
        r.font.size = Pt(9)
        r.font.color.rgb = GRIS
    doc.add_paragraph()

    # ====== Síntesis ejecutiva ======
    h = doc.add_paragraph()
    r = h.add_run("Síntesis ejecutiva")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = AZUL

    cat_dom = max(riesgo.get("categorias", {}).items(), key=lambda x: x[1], default=("—", 0))
    cat_dom_label = cat_dom[0].replace("_", " ").capitalize()

    # Determinar tendencia
    delta = tend.get("delta_vs_avg", 0)
    if delta >= 5:
        trend_phrase = f"al alza (+{delta} vs. promedio semanal {tend['score_avg']})"
        trend_color = ROJO
    elif delta <= -5:
        trend_phrase = f"a la baja ({delta} vs. promedio semanal {tend['score_avg']})"
        trend_color = VERDE
    else:
        trend_phrase = f"estable (Δ{delta:+.1f} vs. promedio semanal {tend['score_avg']})"
        trend_color = GRIS

    p = doc.add_paragraph()
    p.add_run("El nivel de riesgo político se ubica en ")
    rr = p.add_run(f"{score}/100 · {nivel}")
    rr.bold = True
    rr.font.color.rgb = color_nivel(nivel)
    p.add_run(", con tendencia ")
    rr = p.add_run(trend_phrase)
    rr.font.color.rgb = trend_color
    rr.bold = True
    p.add_run(f". Categoría dominante: ")
    rr = p.add_run(cat_dom_label)
    rr.bold = True
    p.add_run(f" ({cat_dom[1]:.0f}/100). En las últimas 24 horas se procesaron ")
    p.add_run(f"{art_24h} ítems de medios y {n_tweets} tweets, ").bold = True
    p.add_run(f"con {len(alertas)} alertas activadas, de las cuales ")
    rr = p.add_run(f"{crit_count} son CRÍTICAS")
    rr.bold = True
    rr.font.color.rgb = ROJO
    p.add_run(f". Sentimiento agregado: {sentimiento}.")

    if matriz:
        p = doc.add_paragraph()
        p.add_run("Factor de mayor presión: ").bold = True
        top_f = matriz[0]
        rr = p.add_run(f"{top_f['nombre']}")
        rr.bold = True
        rr.font.color.rgb = color_nivel(top_f["nivel"])
        p.add_run(f" (Prob {top_f['probabilidad']} · Imp {top_f['impacto']} → score {top_f['score']}).")

    # ====== TENDENCIAS — Top factores P×I ======
    h = doc.add_paragraph()
    r = h.add_run("Análisis de tendencias")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = AZUL

    # Top 5 factores con mini-barras
    p = doc.add_paragraph()
    r = p.add_run("Top 5 factores de riesgo (Probabilidad × Impacto)")
    r.bold = True
    r.font.size = Pt(10)

    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, c in enumerate(["Factor", "Prob/Imp", "Score · Nivel", "Categoría"]):
        hdr[i].text = c
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "1E40AF")

    for f in matriz[:5]:
        row = t.add_row().cells
        row[0].text = f["nombre"]
        # Mini barra texto
        prob_bar = "█" * int(f["probabilidad"] / 10) + "░" * (10 - int(f["probabilidad"] / 10))
        imp_bar = "█" * int(f["impacto"] / 10) + "░" * (10 - int(f["impacto"] / 10))
        row[1].text = f"P {prob_bar} {f['probabilidad']}\nI {imp_bar} {f['impacto']}"
        row[2].text = f"{f['score']} · {f['nivel']}"
        for run in row[2].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = color_nivel(f["nivel"])
        row[3].text = f["categoria"]
        # Tamaño de fuente compacto
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    if not run.bold:
                        run.font.size = Pt(9)

    # ====== Casos persistentes (lo más importante de tendencias) ======
    if tend["alertas_persistentes"]:
        p = doc.add_paragraph()
        p.add_run().add_break()
        r = p.add_run("Casos persistentes · alertas que se repiten")
        r.bold = True
        r.font.size = Pt(10)

        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, c in enumerate(["Caso", "Persistencia", "Categoría"]):
            hdr[i].text = c
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade_cell(hdr[i], "1E40AF")
        for c in tend["alertas_persistentes"][:5]:
            row = t.add_row().cells
            row[0].text = f"[{c['nivel']}] {c['titulo'][:75]}"
            row[1].text = f"{c['dias_activo']} días · {c['ocurrencias']} apariciones"
            row[2].text = c["categoria"]
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
            # destacar nivel
            for run in row[0].paragraphs[0].runs:
                if c["nivel"] == "CRÍTICA":
                    run.font.color.rgb = ROJO

    # Factores que suben/bajan
    if tend["factores_subiendo"] or tend["factores_bajando"]:
        p = doc.add_paragraph()
        r = p.add_run("Movimientos en factores (semana):")
        r.bold = True
        r.font.size = Pt(10)

        if tend["factores_subiendo"]:
            p = doc.add_paragraph()
            r = p.add_run("↑ Subiendo: ")
            r.bold = True
            r.font.color.rgb = ROJO
            r.font.size = Pt(9)
            txt = " · ".join([f"{x['nombre']} ({x['actual']}, +{x['delta']})" for x in tend["factores_subiendo"][:4]])
            run = p.add_run(txt)
            run.font.size = Pt(9)

        if tend["factores_bajando"]:
            p = doc.add_paragraph()
            r = p.add_run("↓ Bajando: ")
            r.bold = True
            r.font.color.rgb = VERDE
            r.font.size = Pt(9)
            txt = " · ".join([f"{x['nombre']} ({x['actual']}, {x['delta']})" for x in tend["factores_bajando"][:3]])
            run = p.add_run(txt)
            run.font.size = Pt(9)

    # ====== ALERTAS CRÍTICAS — Página 3 ======
    doc.add_page_break()
    h = doc.add_paragraph()
    r = h.add_run("Alertas críticas en curso")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = AZUL
    crit = [a for a in alertas if a["nivel"] == "CRÍTICA"][:5]
    if not crit:
        doc.add_paragraph("Sin alertas críticas en la ventana actual.").italic = True
    for a in crit:
        # Title
        p = doc.add_paragraph()
        rr = p.add_run("● ")
        rr.font.color.rgb = ROJO
        rr.font.size = Pt(13)
        rr.bold = True
        rr = p.add_run(a["titulo"])
        rr.bold = True
        rr.font.size = Pt(10)
        # Meta
        p2 = doc.add_paragraph()
        rr = p2.add_run(f"  {a['categoria']} · {a.get('region') or '—'} · {_fmt_dt(a.get('timestamp', ''))} · fuente: {a['fuente']}")
        rr.italic = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = GRIS
        # Resumen
        p3 = doc.add_paragraph()
        rr = p3.add_run(f"  {a.get('resumen', '')[:240]}")
        rr.font.size = Pt(9.5)
        # Acción
        p4 = doc.add_paragraph()
        rr = p4.add_run("  → ACCIÓN: ")
        rr.bold = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = AZUL
        rr = p4.add_run(a.get("accion", ""))
        rr.font.size = Pt(9)
        # URL
        if a.get("url"):
            p5 = doc.add_paragraph()
            rr = p5.add_run(f"  🔗 {a['url']}")
            rr.font.size = Pt(8)
            rr.font.color.rgb = AZUL

    # ====== Footer / Notas ======
    p = doc.add_paragraph()
    p.add_run().add_break()
    p_b = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "8")
    top.set(qn("w:color"), "CBD5E1")
    pBdr.append(top)
    p_b.append(pBdr)

    p = doc.add_paragraph()
    rr = p.add_run(
        "APURISK 1.0 · Plataforma OSINT de Riesgos Políticos del Perú · "
        f"Fuentes: medios peruanos (RPP, La República, El Comercio, Gestión, IDL, Infobae, CNN, Servindi, Caretas, Trome), "
        "Defensoría del Pueblo, Congreso, Twitter/X (API v2), GDELT. "
        f"Snapshot generado: {_fmt_dt_full(fecha_gen)}. "
        f"Análisis de tendencias basado en {tend['snapshots']} snapshots de los últimos 7 días."
    )
    rr.font.size = Pt(7)
    rr.font.color.rgb = GRIS
    rr.italic = True

    doc.save(output_path)
    return output_path


# ============================================================================
# PDF (ReportLab)
# ============================================================================
def generar_ejecutivo_pdf(output_path: str, snapshot: dict, snapshot_dir: str) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, KeepTogether, HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    AZUL = colors.HexColor("#1E40AF")
    AZUL_OSC = colors.HexColor("#0F172A")
    GRIS = colors.HexColor("#64748B")
    GRIS_BG = colors.HexColor("#F8FAFC")
    GRIS_LINEA = colors.HexColor("#CBD5E1")
    ROJO = colors.HexColor("#DC2626")
    NARANJA = colors.HexColor("#EA580C")
    AMARILLO = colors.HexColor("#D48C0B")
    VERDE = colors.HexColor("#16A34A")

    def color_nivel(nivel):
        return {"CRÍTICO": ROJO, "CRÍTICA": ROJO, "ALTO": NARANJA, "ALTA": NARANJA,
                "MEDIO": AMARILLO, "MEDIA": AMARILLO, "BAJO": VERDE, "BAJA": VERDE}.get(nivel, GRIS)

    def esc(s):
        if s is None: return ""
        return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=1.8*cm, rightMargin=1.8*cm,
        topMargin=1.5*cm, bottomMargin=1.5*cm,
        title="APURISK · Reporte Ejecutivo Diario", author="APURISK 1.0",
    )

    ss = getSampleStyleSheet()
    s_title = ParagraphStyle("title", parent=ss["Title"], fontSize=22, textColor=AZUL_OSC,
                              alignment=TA_LEFT, fontName="Helvetica-Bold", spaceAfter=2)
    s_subtitle = ParagraphStyle("subtitle", parent=ss["Normal"], fontSize=12,
                                 textColor=AZUL, fontName="Helvetica-Oblique", spaceAfter=8)
    s_meta = ParagraphStyle("meta", parent=ss["Normal"], fontSize=9, textColor=GRIS, spaceAfter=12)
    s_h1 = ParagraphStyle("h1", parent=ss["Heading1"], fontSize=14, textColor=AZUL,
                           fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=8)
    s_body = ParagraphStyle("body", parent=ss["Normal"], fontSize=9.5, leading=13,
                             textColor=AZUL_OSC, alignment=TA_JUSTIFY)
    s_small = ParagraphStyle("small", parent=ss["Normal"], fontSize=8.5, leading=11, textColor=GRIS)
    s_kpi_label = ParagraphStyle("kpi_label", parent=ss["Normal"], fontSize=8,
                                  textColor=GRIS, alignment=TA_CENTER, fontName="Helvetica-Bold")
    s_kpi_value = ParagraphStyle("kpi_value", parent=ss["Normal"], fontSize=28,
                                  alignment=TA_CENTER, fontName="Helvetica-Bold")
    s_kpi_sub = ParagraphStyle("kpi_sub", parent=ss["Normal"], fontSize=9,
                                textColor=GRIS, alignment=TA_CENTER)
    s_alert_title = ParagraphStyle("alert_title", parent=ss["Normal"], fontSize=10,
                                    textColor=AZUL_OSC, fontName="Helvetica-Bold", spaceBefore=8)
    s_url = ParagraphStyle("url", parent=ss["Normal"], fontSize=7.5, textColor=AZUL,
                            leading=10, leftIndent=12)

    story = []

    # Datos
    riesgo = snapshot.get("riesgo", {})
    score = riesgo.get("global", 0)
    nivel = riesgo.get("nivel", "—")
    alertas = snapshot.get("alertas", [])
    crit = [a for a in alertas if a["nivel"] == "CRÍTICA"]
    matriz = snapshot.get("matriz_riesgo", [])
    fecha_gen = snapshot.get("generado", "")
    art_24h = snapshot.get("n_articulos_24h", 0)
    n_tweets = snapshot.get("n_tweets", 0)
    n_conf = snapshot.get("n_conflictos", 0)
    sentimiento = riesgo.get("sentimiento_promedio", 0)
    cat_dom = max(riesgo.get("categorias", {}).items(), key=lambda x: x[1], default=("—", 0))
    cat_dom_label = cat_dom[0].replace("_", " ").capitalize()

    tend = _aggregate_tendencias(snapshot_dir)

    # ====== HEADER ======
    story.append(Paragraph(
        f'<font color="#0F172A">APURISK 1.0</font> '
        f'<font color="#1E40AF" size="14"> · Reporte Ejecutivo Diario</font>',
        s_title
    ))
    story.append(Paragraph(f"Riesgo Político · Perú · {esc(_fmt_dt_full(fecha_gen))}", s_meta))
    story.append(HRFlowable(width="100%", thickness=1.5, color=AZUL, spaceBefore=2, spaceAfter=10))

    # ====== KPIs ======
    cn = color_nivel(nivel)
    kpi_data = [[
        Paragraph("SCORE GLOBAL", s_kpi_label),
        Paragraph("ALERTAS CRÍTICAS", s_kpi_label),
        Paragraph("CONFLICTOS SOC.", s_kpi_label),
        Paragraph("COBERTURA 24H", s_kpi_label),
    ], [
        Paragraph(f'<font color="{cn.hexval()}">{score}</font>', s_kpi_value),
        Paragraph(f'<font color="{ROJO.hexval()}">{len(crit)}</font>', s_kpi_value),
        Paragraph(f'<font color="{NARANJA.hexval()}">{n_conf}</font>', s_kpi_value),
        Paragraph(f'<font color="{AZUL.hexval()}">{art_24h + n_tweets}</font>', s_kpi_value),
    ], [
        Paragraph(f'<font color="{cn.hexval()}">{nivel}</font>', s_kpi_sub),
        Paragraph(f"de {len(alertas)} totales", s_kpi_sub),
        Paragraph("activos", s_kpi_sub),
        Paragraph(f"{art_24h} medios + {n_tweets} X", s_kpi_sub),
    ]]
    kpi_table = Table(kpi_data, colWidths=[4.4*cm]*4, rowHeights=[0.7*cm, 1.4*cm, 0.6*cm])
    kpi_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), GRIS_BG),
        ("BOX", (0, 0), (-1, -1), 0.5, GRIS_LINEA),
        ("LINEAFTER", (0, 0), (-2, -1), 0.5, GRIS_LINEA),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(kpi_table)
    story.append(Spacer(1, 14))

    # ====== Síntesis ejecutiva ======
    story.append(Paragraph("Síntesis ejecutiva", s_h1))

    delta = tend.get("delta_vs_avg", 0)
    if delta >= 5:
        trend_phrase = f'<font color="{ROJO.hexval()}"><b>al alza (+{delta} vs. promedio semanal {tend["score_avg"]})</b></font>'
    elif delta <= -5:
        trend_phrase = f'<font color="{VERDE.hexval()}"><b>a la baja ({delta} vs. promedio semanal {tend["score_avg"]})</b></font>'
    else:
        trend_phrase = f'<b>estable (Δ{delta:+.1f} vs. promedio semanal {tend["score_avg"]})</b>'

    sintesis = (
        f'El nivel de riesgo político se ubica en '
        f'<b><font color="{cn.hexval()}">{score}/100 · {nivel}</font></b>, '
        f'con tendencia {trend_phrase}. '
        f'Categoría dominante: <b>{esc(cat_dom_label)}</b> ({cat_dom[1]:.0f}/100). '
        f'En las últimas 24 horas se procesaron <b>{art_24h} ítems de medios</b> '
        f'y <b>{n_tweets} tweets</b>, con {len(alertas)} alertas activadas, '
        f'de las cuales <b><font color="{ROJO.hexval()}">{len(crit)} son CRÍTICAS</font></b>. '
        f'Sentimiento agregado: <b>{sentimiento}</b>.'
    )
    story.append(Paragraph(sintesis, s_body))
    if matriz:
        top_f = matriz[0]
        c_top = color_nivel(top_f["nivel"])
        story.append(Paragraph(
            f'<b>Factor de mayor presión:</b> '
            f'<b><font color="{c_top.hexval()}">{esc(top_f["nombre"])}</font></b> '
            f'(Prob {top_f["probabilidad"]} · Imp {top_f["impacto"]} → score <b>{top_f["score"]}</b>).',
            s_body
        ))
    story.append(Spacer(1, 8))

    # ====== ANÁLISIS DE TENDENCIAS — TOP FACTORES ======
    story.append(Paragraph("Análisis de tendencias", s_h1))
    story.append(Paragraph("<b>Top 5 factores de riesgo (Probabilidad × Impacto)</b>", s_body))

    # Tabla de factores con barras visuales en Paragraph
    rows = [["Factor", "Probabilidad", "Impacto", "Score · Nivel"]]
    for f in matriz[:5]:
        prob_pct = int(f["probabilidad"])
        imp_pct = int(f["impacto"])
        bar_p = f"█" * int(prob_pct / 10) + "░" * (10 - int(prob_pct / 10))
        bar_i = f"█" * int(imp_pct / 10) + "░" * (10 - int(imp_pct / 10))
        c_n = color_nivel(f["nivel"])
        rows.append([
            Paragraph(f'<b>{esc(f["nombre"])}</b><br/>'
                       f'<font color="{GRIS.hexval()}" size="7">{esc(f["categoria"])}</font>', s_small),
            Paragraph(f'<font name="Courier" size="7">{bar_p}</font> <b>{prob_pct}</b>', s_small),
            Paragraph(f'<font name="Courier" size="7">{bar_i}</font> <b>{imp_pct}</b>', s_small),
            Paragraph(f'<b><font color="{c_n.hexval()}" size="11">{f["score"]}</font></b><br/>'
                       f'<font color="{c_n.hexval()}" size="8"><b>{f["nivel"]}</b></font>', s_small),
        ])
    t = Table(rows, colWidths=[6.5*cm, 4.5*cm, 4.5*cm, 2.5*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), AZUL),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, GRIS_BG]),
        ("BOX", (0, 0), (-1, -1), 0.4, GRIS_LINEA),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, GRIS_LINEA),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(t)
    story.append(Spacer(1, 10))

    # ====== Casos persistentes ======
    if tend["alertas_persistentes"]:
        story.append(Paragraph(
            f'<b>Casos persistentes</b> · alertas detectadas en múltiples días '
            f'<font color="{GRIS.hexval()}" size="8">({tend["snapshots"]} snapshots)</font>',
            s_body
        ))
        rows = [["Caso", "Persistencia", "Categoría"]]
        for c in tend["alertas_persistentes"][:5]:
            c_n = color_nivel(c["nivel"])
            rows.append([
                Paragraph(f'<font color="{c_n.hexval()}"><b>[{c["nivel"]}]</b></font> '
                           f'{esc(c["titulo"][:75])}', s_small),
                Paragraph(f'<b>{c["dias_activo"]}</b> días · {c["ocurrencias"]} apariciones', s_small),
                Paragraph(esc(c["categoria"]), s_small),
            ])
        t = Table(rows, colWidths=[10*cm, 4*cm, 4*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), AZUL),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 8.5),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, GRIS_BG]),
            ("BOX", (0, 0), (-1, -1), 0.4, GRIS_LINEA),
            ("INNERGRID", (0, 0), (-1, -1), 0.3, GRIS_LINEA),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 8))

    # ====== Movimientos en factores ======
    if tend["factores_subiendo"] or tend["factores_bajando"]:
        if tend["factores_subiendo"]:
            txt = f'<font color="{ROJO.hexval()}"><b>↑ Subiendo:</b></font> ' + " · ".join(
                [f'{esc(x["nombre"])} ({x["actual"]}, +{x["delta"]})' for x in tend["factores_subiendo"][:4]]
            )
            story.append(Paragraph(txt, s_body))
        if tend["factores_bajando"]:
            txt = f'<font color="{VERDE.hexval()}"><b>↓ Bajando:</b></font> ' + " · ".join(
                [f'{esc(x["nombre"])} ({x["actual"]}, {x["delta"]})' for x in tend["factores_bajando"][:3]]
            )
            story.append(Paragraph(txt, s_body))

    # ====== ALERTAS CRÍTICAS — Página 3 ======
    story.append(PageBreak())
    story.append(Paragraph("Alertas críticas en curso", s_h1))
    if not crit:
        story.append(Paragraph("<i>Sin alertas críticas en la ventana actual.</i>", s_body))
    for a in crit[:5]:
        block = []
        block.append(Paragraph(
            f'<font color="{ROJO.hexval()}">●</font> <b>{esc(a["titulo"])}</b>',
            s_alert_title
        ))
        block.append(Paragraph(
            f'<i><font color="{GRIS.hexval()}">'
            f'{esc(a["categoria"])} · {esc(a.get("region") or "—")} · '
            f'{esc(_fmt_dt(a.get("timestamp", "")))} · fuente: {esc(a["fuente"])}</font></i>',
            s_small
        ))
        block.append(Paragraph(esc(a.get("resumen", "")[:240]), s_body))
        block.append(Paragraph(
            f'<font color="{AZUL.hexval()}"><b>→ ACCIÓN:</b></font> {esc(a.get("accion", ""))}',
            s_small
        ))
        if a.get("url"):
            block.append(Paragraph(
                f'🔗 <a href="{a["url"]}" color="{AZUL.hexval()}"><u>{esc(a["url"][:90])}</u></a>',
                s_url
            ))
        block.append(Spacer(1, 4))
        story.append(KeepTogether(block))

    # Footer
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=0.5, color=GRIS_LINEA, spaceBefore=4, spaceAfter=4))
    story.append(Paragraph(
        f'APURISK 1.0 · Plataforma OSINT de Riesgos Políticos del Perú · '
        f'Fuentes: medios peruanos (RPP, La República, El Comercio, Gestión, Infobae, Servindi, Caretas, Trome), '
        f'Defensoría del Pueblo, Congreso, Twitter/X (API v2), GDELT. '
        f'Snapshot generado: {esc(_fmt_dt_full(fecha_gen))}. '
        f'Análisis de tendencias basado en {tend["snapshots"]} snapshots de los últimos 7 días.',
        ParagraphStyle("footer", parent=ss["Normal"], fontSize=7,
                        textColor=GRIS, fontName="Helvetica-Oblique", alignment=TA_JUSTIFY)
    ))

    doc.build(story)
    return output_path
