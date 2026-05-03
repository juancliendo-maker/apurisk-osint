"""Reporte 24h — síntesis ejecutiva imprimible.

Genera HTML y DOCX con la actividad de las últimas 24 horas, ideal para
briefing matutino del equipo de consultoría política.
"""
from __future__ import annotations
from datetime import datetime
import html as _html
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _esc(s: str | None) -> str:
    return _html.escape(s or "", quote=True)


def _fmt_h(h):
    if h is None or h == float("inf"):
        return "—"
    if h < 1:
        return f"hace {int(h*60)}m"
    return f"hace {h:.1f}h"


def generar_reporte_24h_html(output_path: str, art_24, conf_24, alertas_24, riesgo, matriz, modo: str):
    top_factores = matriz[:5]
    headlines = sorted(art_24, key=lambda a: a.hours_ago())

    factores_html = ""
    for f in top_factores:
        factores_html += f"""
        <tr>
          <td><strong>{_esc(f['nombre'])}</strong><br><span class='sub'>{_esc(f['categoria'])}</span></td>
          <td style='text-align:center;'>{f['probabilidad']}</td>
          <td style='text-align:center;'>{f['impacto']}</td>
          <td style='text-align:center;'><strong class='nivel-{f['nivel']}'>{f['score']}</strong></td>
          <td style='text-align:center;'>{f['tendencia']}</td>
        </tr>
        """

    alertas_html = ""
    for a in alertas_24[:20]:
        alertas_html += f"""
        <div class='alerta {a['nivel']}'>
          <div class='head'><span class='tag'>{a['nivel']}</span> <strong>{_esc(a['titulo'])}</strong></div>
          <div class='resumen'>{_esc(a['resumen'])}</div>
          <div class='meta'>{_esc(a['categoria'])} · {_esc(a['fuente'])} · {_fmt_h(a['hours_ago'])}
            {f"· <a href='{_esc(a['url'])}' target='_blank'>fuente</a>" if a.get('url') else ''}
          </div>
        </div>
        """

    headlines_html = ""
    for a in headlines[:30]:
        url = _esc(a.url or "#")
        headlines_html += f"""
        <div class='headline'>
          <span class='src'>{_esc(a.source_name)}</span> · <span class='time'>{_fmt_h(a.hours_ago())}</span>
          <div class='title'><a href='{url}' target='_blank'>{_esc(a.title)}</a></div>
          <div class='sum'>{_esc((a.summary or '')[:300])}</div>
        </div>
        """

    html = f"""<!DOCTYPE html>
<html lang='es'><head><meta charset='utf-8'/>
<title>APURISK · Reporte 24h · {datetime.now().strftime('%d %b %Y')}</title>
<style>
  body {{ font-family: 'Inter', -apple-system, sans-serif; background:#fff; color:#0f172a; margin:0; padding: 40px; max-width: 880px; margin: 0 auto; }}
  h1 {{ font-size: 26px; margin-bottom: 4px; color:#0f172a; }}
  h2 {{ font-size: 18px; margin: 28px 0 12px; color:#1e40af; border-bottom: 2px solid #1e40af; padding-bottom: 6px;}}
  .sub {{ color:#64748b; font-size: 12px;}}
  .meta-top {{ color:#64748b; font-size: 13px; border-bottom: 1px solid #e2e8f0; padding-bottom: 12px; margin-bottom: 24px;}}
  .resumen-box {{ background:#f1f5f9; border-left: 4px solid #1e40af; padding: 14px 18px; border-radius: 6px; margin-bottom: 18px; line-height: 1.7;}}
  .kpi-row {{ display:flex; gap: 14px; margin: 14px 0;}}
  .kpi {{ background: #f8fafc; padding: 12px 16px; border-radius: 6px; border: 1px solid #e2e8f0; flex:1;}}
  .kpi .lbl {{ font-size: 10px; text-transform: uppercase; color:#64748b; letter-spacing:.5px;}}
  .kpi .val {{ font-size: 24px; font-weight: 700; margin-top: 4px;}}
  table {{ width:100%; border-collapse: collapse; margin-bottom: 14px;}}
  th, td {{ text-align:left; padding: 8px; border-bottom: 1px solid #e2e8f0; font-size: 13px; }}
  th {{ background:#f1f5f9; font-weight:600; text-transform:uppercase; font-size:11px; color:#64748b;}}
  .nivel-CRÍTICO {{color:#b91c1c;}} .nivel-ALTO {{color:#c2410c;}} .nivel-MEDIO {{color:#a16207;}} .nivel-BAJO {{color:#15803d;}}
  .alerta {{ border-left: 3px solid #f59e0b; padding: 10px 14px; margin-bottom: 8px; background: #fffbeb; border-radius: 4px;}}
  .alerta.CRÍTICA {{ border-left-color: #dc2626; background: #fef2f2;}}
  .alerta.ALTA {{ border-left-color: #ea580c; background: #fff7ed;}}
  .alerta .tag {{ background:#dc2626; color:white; padding: 2px 8px; border-radius:3px; font-size:10px; font-weight:700; letter-spacing:1px;}}
  .alerta.ALTA .tag {{ background:#ea580c;}}
  .alerta.MEDIA .tag {{ background:#f59e0b;}}
  .alerta .resumen {{ color:#475569; font-size:12.5px; margin-top: 4px;}}
  .alerta .meta {{ color:#64748b; font-size: 11px; margin-top: 4px;}}
  .alerta a {{ color:#1e40af;}}
  .headline {{ padding: 10px 0; border-bottom: 1px solid #e2e8f0;}}
  .headline .src {{ font-size: 11px; color:#1e40af; font-weight:600; text-transform: uppercase;}}
  .headline .time {{ font-size: 11px; color:#94a3b8;}}
  .headline .title a {{ color:#0f172a; text-decoration: none; font-weight:600; font-size:14px;}}
  .headline .title a:hover {{ color:#1e40af;}}
  .headline .sum {{ color:#475569; font-size:12.5px; margin-top: 4px;}}
  @media print {{ body {{ padding: 24px;}} a {{ color:#1e40af;}} }}
</style></head>
<body>
<h1>APURISK · Reporte de Riesgo Político · Últimas 24 horas</h1>
<div class='meta-top'>Generado: {datetime.now().strftime('%d %b %Y · %H:%M')} · Modo: {modo}</div>

<div class='resumen-box'>
  En las últimas <strong>24 horas</strong>: <strong>{len(art_24)}</strong> piezas de cobertura procesadas,
  <strong>{len(alertas_24)}</strong> alertas activadas
  ({len([a for a in alertas_24 if a['nivel']=='CRÍTICA'])} críticas).
  Score global de riesgo: <strong class='nivel-{riesgo['nivel']}'>{riesgo['global']}/100 · {riesgo['nivel']}</strong>.
  Categoría dominante: <strong>{max(riesgo['categorias'].items(), key=lambda x: x[1])[0].replace('_',' ')}</strong>
  ({max(riesgo['categorias'].values()):.0f}/100).
</div>

<div class='kpi-row'>
  <div class='kpi'><div class='lbl'>Score global</div><div class='val nivel-{riesgo['nivel']}'>{riesgo['global']}</div></div>
  <div class='kpi'><div class='lbl'>Alertas críticas</div><div class='val nivel-CRÍTICO'>{len([a for a in alertas_24 if a['nivel']=='CRÍTICA'])}</div></div>
  <div class='kpi'><div class='lbl'>Conflictos 24h</div><div class='val'>{len(conf_24)}</div></div>
  <div class='kpi'><div class='lbl'>Sentimiento</div><div class='val'>{riesgo['sentimiento_promedio']}</div></div>
</div>

<h2>Top 5 factores de riesgo</h2>
<table>
  <tr><th>Factor</th><th>Prob</th><th>Imp</th><th>Score</th><th>Tend</th></tr>
  {factores_html}
</table>

<h2>Alertas activas en la ventana</h2>
{alertas_html or '<em>Sin alertas activas en las últimas 24 horas.</em>'}

<h2>Cobertura procesada — Headlines</h2>
{headlines_html or '<em>Sin cobertura en las últimas 24 horas.</em>'}

</body></html>
"""
    Path(output_path).write_text(html, encoding="utf-8")
    return output_path


def generar_reporte_24h_docx(output_path: str, art_24, conf_24, alertas_24, riesgo, matriz, modo: str):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("APURISK · Reporte 24h")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Riesgo Político · Perú\nGenerado: {datetime.now().strftime('%d %b %Y · %H:%M')} · Modo: {modo}\n")

    doc.add_heading("Síntesis ejecutiva", level=1)
    p = doc.add_paragraph()
    p.add_run(
        f"En las últimas 24 horas se procesaron {len(art_24)} piezas de cobertura, "
        f"se activaron {len(alertas_24)} alertas "
        f"({len([a for a in alertas_24 if a['nivel']=='CRÍTICA'])} críticas) y se monitorearon "
        f"{len(conf_24)} eventos de conflictividad social. "
    )
    p.add_run("Score global: ").bold = True
    rr = p.add_run(f"{riesgo['global']}/100 · {riesgo['nivel']}. ")
    rr.bold = True
    p.add_run(
        f"Categoría dominante: {max(riesgo['categorias'].items(), key=lambda x: x[1])[0].replace('_',' ')} "
        f"({max(riesgo['categorias'].values()):.0f}/100)."
    )

    doc.add_heading("Top 5 factores de riesgo", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Factor", "Prob", "Imp", "Score", "Tend"]):
        hdr[i].text = h
    for f in matriz[:5]:
        r = t.add_row().cells
        r[0].text = f"{f['nombre']} ({f['categoria']})"
        r[1].text = str(f["probabilidad"])
        r[2].text = str(f["impacto"])
        r[3].text = f"{f['score']} · {f['nivel']}"
        r[4].text = f["tendencia"]

    doc.add_heading("Alertas activas", level=1)
    if not alertas_24:
        doc.add_paragraph("Sin alertas activas en las últimas 24 horas.")
    for a in alertas_24:
        p = doc.add_paragraph()
        rr = p.add_run(f"[{a['nivel']}] ")
        rr.bold = True
        rr.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C) if a["nivel"] == "CRÍTICA" else RGBColor(0xC2, 0x41, 0x0C)
        p.add_run(a["titulo"]).bold = True
        p.add_run(f"\n{a['resumen']}\n")
        meta = doc.add_paragraph()
        meta.add_run(f"Categoría: {a['categoria']} · Fuente: {a['fuente']} · {_fmt_h(a['hours_ago'])}").italic = True
        if a.get("url"):
            meta.add_run(f"\nURL: {a['url']}")

    doc.add_heading("Headlines · cobertura procesada", level=1)
    for a in sorted(art_24, key=lambda x: x.hours_ago())[:30]:
        p = doc.add_paragraph()
        p.add_run(f"[{a.source_name} · {_fmt_h(a.hours_ago())}] ").italic = True
        p.add_run(a.title).bold = True
        if a.summary:
            doc.add_paragraph((a.summary or "")[:280])
        if a.url:
            url_p = doc.add_paragraph()
            url_p.add_run(f"URL: {a.url}").italic = True

    doc.save(output_path)
    return output_path
