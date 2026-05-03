"""Reporte de Alertas Inmediatas — feed crítico con triggers.

HTML imprimible y DOCX para distribución por email/Slack a stakeholders.
"""
from __future__ import annotations
from datetime import datetime
import html as _html
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _esc(s: str | None) -> str:
    return _html.escape(s or "", quote=True)


def _fmt_h(h):
    if h is None or h == float("inf"):
        return "—"
    if h < 1:
        return f"hace {int(h*60)}m"
    if h < 24:
        return f"hace {h:.1f}h"
    return f"hace {h/24:.1f}d"


_MESES = {1:"ene",2:"feb",3:"mar",4:"abr",5:"may",6:"jun",7:"jul",8:"ago",9:"set",10:"oct",11:"nov",12:"dic"}


def _fmt_dt(iso_str):
    if not iso_str:
        return "—"
    try:
        from datetime import datetime
        dt = datetime.fromisoformat(iso_str)
        return f"{dt.day:02d} {_MESES.get(dt.month, '?')} {dt.strftime('%H:%M')}"
    except Exception:
        return iso_str[:16]


def _short_url(url, max_len=55):
    if not url: return ""
    import re
    m = re.match(r"https?://([^/]+)(/.*)?", url)
    if not m: return url[:max_len] + ("…" if len(url) > max_len else "")
    host = m.group(1); path = m.group(2) or ""
    if len(path) > max_len - len(host):
        path = path[:max_len - len(host) - 1] + "…"
    return f"{host}{path}"


def generar_alertas_html(output_path: str, alertas, modo: str):
    crit = [a for a in alertas if a["nivel"] == "CRÍTICA"]
    alta = [a for a in alertas if a["nivel"] == "ALTA"]

    def alerta_block(a):
        url_link = (
            f"<a href='{_esc(a['url'])}' target='_blank' rel='noopener' title='{_esc(a['url'])}'>"
            f"🔗 {_esc(_short_url(a['url']))}</a>"
            if a.get("url") else ""
        )
        region = f"<span class='region'>📍 {_esc(a['region'])}</span>" if a.get("region") else ""
        ts = a.get("timestamp", "")
        titulo_html = _esc(a['titulo'])
        if a.get('url'):
            titulo_html = f"<a href='{_esc(a['url'])}' target='_blank' rel='noopener'>{titulo_html}</a>"
        return f"""
        <div class='alerta {a['nivel']}'>
          <div class='head'>
            <span class='tag'>{a['nivel']}</span>
            <span class='cat'>{_esc(a['categoria'])}</span>
            {region}
            <span class='time' title='{_esc(ts)}'>🕒 {_fmt_dt(ts)} · {_fmt_h(a['hours_ago'])}</span>
          </div>
          <div class='titulo'>{titulo_html}</div>
          <div class='resumen'>{_esc(a['resumen'])}</div>
          <div class='accion'><strong>ACCIÓN RECOMENDADA</strong> {_esc(a['accion'])}</div>
          <div class='links'>
            <span style='color:#94a3b8; font-size:11px;'>Fuente: {_esc(a['fuente'])}</span>
            {url_link}
            <span class='regla'>Regla: {_esc(a['regla'])}</span>
          </div>
        </div>
        """

    crit_html = "".join(alerta_block(a) for a in crit) or "<em style='color:#64748b;'>Sin alertas críticas en este momento.</em>"
    alta_html = "".join(alerta_block(a) for a in alta) or "<em style='color:#64748b;'>Sin alertas de nivel alto.</em>"

    html = f"""<!DOCTYPE html>
<html lang='es'><head><meta charset='utf-8'/>
<title>APURISK · Alertas Inmediatas · {datetime.now().strftime('%d %b %Y · %H:%M')}</title>
<style>
  body {{ font-family: 'Inter', -apple-system, sans-serif; background:#0a0e1a; color:#f1f5f9; margin:0; padding: 32px; max-width: 880px; margin: 0 auto;}}
  h1 {{ color:#fff; font-size: 26px; margin-bottom: 4px;}}
  h2 {{ color:#fff; font-size: 18px; margin: 24px 0 12px; padding: 8px 14px; background: linear-gradient(90deg, #dc2626, transparent); border-radius: 4px;}}
  h2.alta {{ background: linear-gradient(90deg, #ea580c, transparent);}}
  .meta-top {{ color:#94a3b8; font-size: 13px; padding-bottom: 12px; border-bottom: 1px solid #1e293b; margin-bottom: 18px;}}
  .stats {{ display:flex; gap:14px; margin-bottom: 20px;}}
  .stat {{ background:#1e293b; padding: 14px 18px; border-radius: 8px; flex:1;}}
  .stat .lbl {{ font-size:11px; color:#94a3b8; text-transform: uppercase; letter-spacing:1px;}}
  .stat .val {{ font-size: 28px; font-weight:700; margin-top: 4px;}}
  .stat.crit .val {{color: #ef4444;}}
  .stat.alta .val {{color: #f97316;}}
  .alerta {{ background: #1e293b; border-left: 4px solid #f97316; padding: 16px; border-radius: 6px; margin-bottom: 12px;}}
  .alerta.CRÍTICA {{ border-left-color: #ef4444; background: linear-gradient(90deg, rgba(239,68,68,0.1), #1e293b);}}
  .alerta.ALTA {{ border-left-color: #f97316;}}
  .alerta .head {{ display:flex; gap: 10px; align-items: center; margin-bottom: 8px; flex-wrap: wrap;}}
  .alerta .tag {{ font-size: 10px; font-weight: 700; padding: 3px 10px; border-radius: 4px; letter-spacing:1.5px;}}
  .alerta.CRÍTICA .tag {{ background:#ef4444; color:#fff;}}
  .alerta.ALTA .tag {{ background:#f97316; color:#fff;}}
  .alerta .cat {{ background:#334155; padding: 2px 8px; border-radius: 4px; font-size:11px;}}
  .alerta .region {{ color:#cbd5e1; font-size: 12px;}}
  .alerta .time {{ color:#94a3b8; font-size: 12px; margin-left: auto;}}
  .alerta .titulo {{ font-size: 15px; font-weight:600; margin-bottom: 6px; color:#fff;}}
  .alerta .resumen {{ font-size: 13px; color:#cbd5e1; line-height: 1.5;}}
  .alerta .accion {{ background: rgba(56,189,248,0.1); border-left: 2px solid #38bdf8; padding: 8px 12px; margin-top: 10px; border-radius: 4px; font-size: 12.5px;}}
  .alerta .accion strong {{ color: #38bdf8; letter-spacing:.5px; font-size: 10.5px; display:block; margin-bottom: 2px;}}
  .alerta .links {{ margin-top: 8px; display:flex; gap:14px; align-items: center; flex-wrap: wrap;}}
  .alerta .links a {{ color:#38bdf8; font-size:12px; text-decoration:none; padding: 2px 8px; background: rgba(56,189,248,0.1); border-radius: 4px;}}
  .alerta .links a:hover {{ background: #38bdf8; color: #0a0e1a;}}
  .alerta .links .regla {{ color:#64748b; font-size: 11px; font-family: monospace;}}
  @media print {{ body {{ background:#fff; color:#0f172a;}} .alerta {{ background:#f8fafc; color:#0f172a;}} .alerta .titulo {{color:#0f172a;}} h1, h2 {{color:#0f172a; background:#f1f5f9;}} }}
</style></head>
<body>

<h1>🚨 APURISK · Alertas Inmediatas</h1>
<div class='meta-top'>Generado: {datetime.now().strftime('%d %b %Y · %H:%M')} · Modo: {modo} · Total alertas: {len(alertas)}</div>

<div class='stats'>
  <div class='stat crit'><div class='lbl'>Críticas</div><div class='val'>{len(crit)}</div></div>
  <div class='stat alta'><div class='lbl'>Altas</div><div class='val'>{len(alta)}</div></div>
  <div class='stat'><div class='lbl'>Total activas</div><div class='val'>{len(alertas)}</div></div>
</div>

<h2>Alertas críticas — acción inmediata</h2>
{crit_html}

<h2 class='alta'>Alertas de nivel alto</h2>
{alta_html}

</body></html>
"""
    Path(output_path).write_text(html, encoding="utf-8")
    return output_path


def generar_alertas_docx(output_path: str, alertas, modo: str):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("APURISK · Alertas Inmediatas")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"Generado: {datetime.now().strftime('%d %b %Y · %H:%M')} · Modo: {modo}\n"
        f"Total alertas: {len(alertas)} "
        f"({len([a for a in alertas if a['nivel']=='CRÍTICA'])} críticas, "
        f"{len([a for a in alertas if a['nivel']=='ALTA'])} altas)\n"
    )

    grupos = [
        ("CRÍTICA", "Alertas críticas — acción inmediata", RGBColor(0xC0, 0x39, 0x2B)),
        ("ALTA", "Alertas de nivel alto", RGBColor(0xC2, 0x41, 0x0C)),
        ("MEDIA", "Alertas de nivel medio", RGBColor(0xA1, 0x62, 0x07)),
    ]
    for nivel, titulo, color in grupos:
        items = [a for a in alertas if a["nivel"] == nivel]
        if not items:
            continue
        h = doc.add_heading(titulo, level=1)
        for a in items:
            p = doc.add_paragraph()
            rr = p.add_run(f"[{a['nivel']}] ")
            rr.bold = True
            rr.font.color.rgb = color
            p.add_run(a["titulo"]).bold = True
            doc.add_paragraph(a["resumen"])
            meta = doc.add_paragraph()
            meta.add_run(
                f"Categoría: {a['categoria']} · Región: {a.get('region') or '—'} · "
                f"Fuente: {a['fuente']} · {_fmt_h(a['hours_ago'])}"
            ).italic = True
            accion = doc.add_paragraph()
            accion.add_run("Acción recomendada: ").bold = True
            accion.add_run(a["accion"])
            if a.get("url"):
                url_p = doc.add_paragraph()
                url_p.add_run(f"URL: {a['url']}").italic = True
            doc.add_paragraph()  # separador

    doc.save(output_path)
    return output_path
