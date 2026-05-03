"""Reporte ejecutivo en Word (.docx) usando python-docx."""
from __future__ import annotations
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


COLOR_ROJO = RGBColor(0xC0, 0x39, 0x2B)
COLOR_AMARILLO = RGBColor(0xD4, 0x8C, 0x0B)
COLOR_VERDE = RGBColor(0x2E, 0x86, 0x32)
COLOR_AZUL = RGBColor(0x1F, 0x4E, 0x79)


def _shade(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _color_nivel(nivel: str):
    return {"ALTO": COLOR_ROJO, "MEDIO": COLOR_AMARILLO, "BAJO": COLOR_VERDE}.get(nivel, COLOR_AZUL)


def generar_reporte_docx(
    output_path: str,
    articulos,
    conflictos,
    proyectos,
    entidades: dict,
    temas: dict,
    riesgo: dict,
    modo: str = "demo",
    ventana: int = 7,
):
    doc = Document()

    # Estilos por defecto
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("APURISK 1.0")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plataforma OSINT de Riesgos Políticos del Perú")
    run.italic = True
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"\nReporte de Inteligencia · {datetime.now().strftime('%d de %B de %Y · %H:%M')}\n")
    p.add_run(f"Ventana de análisis: últimos {ventana} días · Modo: {modo}\n")

    doc.add_paragraph()

    # Resumen ejecutivo
    doc.add_heading("1. Resumen ejecutivo", level=1)

    p = doc.add_paragraph()
    p.add_run("Score de riesgo político global: ").bold = True
    r = p.add_run(f"{riesgo['global']} / 100  →  {riesgo['nivel']}")
    r.bold = True
    r.font.color.rgb = _color_nivel(riesgo["nivel"])

    p = doc.add_paragraph()
    p.add_run("Sentimiento promedio de cobertura: ").bold = True
    p.add_run(f"{riesgo['sentimiento_promedio']} (rango -1 muy negativo a +1 muy positivo).")

    # Hallazgos clave (auto-generado)
    doc.add_heading("Hallazgos clave", level=2)
    hallazgos = _generar_hallazgos(temas, conflictos, entidades, riesgo)
    for h in hallazgos:
        doc.add_paragraph(h, style="List Bullet")

    # Score por categoría
    doc.add_heading("2. Indicadores por categoría", level=1)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Categoría"
    hdr[1].text = "Score (0-100)"
    hdr[2].text = "Lectura"
    for cat, sc in riesgo["categorias"].items():
        row = t.add_row().cells
        row[0].text = cat.replace("_", " ").capitalize()
        row[1].text = str(sc)
        if sc >= 70:
            row[2].text = "Alta atención"
            _shade(row[2], "F8D7DA")
        elif sc >= 45:
            row[2].text = "Vigilancia"
            _shade(row[2], "FFF3CD")
        else:
            row[2].text = "Estable"
            _shade(row[2], "D4EDDA")

    # Temas relevantes
    doc.add_heading("3. Temas con mayor cobertura", level=1)
    conteos = sorted(temas.get("conteos", {}).items(), key=lambda x: -x[1])
    if conteos:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light List Accent 1"
        h = t.rows[0].cells
        h[0].text = "Tema"
        h[1].text = "Menciones"
        for k, v in conteos:
            r = t.add_row().cells
            r[0].text = k.replace("_", " ").capitalize()
            r[1].text = str(v)

    # Conflictos sociales
    doc.add_heading("4. Conflictos sociales en seguimiento", level=1)
    if conflictos:
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        h = t.rows[0].cells
        for i, txt in enumerate(["Región", "Tipo", "Severidad", "Estado"]):
            h[i].text = txt
        for c in conflictos:
            raw = c.raw or {}
            r = t.add_row().cells
            r[0].text = raw.get("region") or c.region or "—"
            r[1].text = raw.get("tipo", "—")
            sev = raw.get("severidad", "—")
            r[2].text = sev
            if sev == "alta":
                _shade(r[2], "F8D7DA")
            elif sev == "media":
                _shade(r[2], "FFF3CD")
            r[3].text = raw.get("estado", "—")
        doc.add_paragraph()
        for c in conflictos:
            doc.add_paragraph(f"• {c.title} — {c.summary}", style="Normal")

    # Proyectos de Ley
    doc.add_heading("5. Proyectos de ley en seguimiento", level=1)
    if proyectos:
        for p_ in proyectos:
            estado = (p_.raw or {}).get("estado", "—")
            cat = (p_.raw or {}).get("categoria", "—")
            par = doc.add_paragraph()
            par.add_run(p_.title).bold = True
            par.add_run(f"  ·  Estado: {estado} · Categoría: {cat}\n")
            par.add_run(p_.summary)

    # Top entidades
    doc.add_heading("6. Top entidades mencionadas", level=1)
    sub = [
        ("Instituciones", entidades.get("instituciones", [])),
        ("Partidos políticos", entidades.get("partidos", [])),
        ("Empresas en zona de riesgo", entidades.get("empresas_riesgo", [])),
        ("Regiones", entidades.get("regiones", [])),
    ]
    for nombre, items in sub:
        if not items:
            continue
        doc.add_heading(nombre, level=2)
        line = ", ".join([f"{k} ({v})" for k, v in items])
        doc.add_paragraph(line)

    # Anexo: cobertura
    doc.add_heading("7. Anexo — cobertura procesada", level=1)
    doc.add_paragraph(f"Total de artículos procesados: {len(articulos)}.")
    fuentes = {}
    for a in articulos:
        fuentes[a.source_name] = fuentes.get(a.source_name, 0) + 1
    if fuentes:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light List Accent 1"
        h = t.rows[0].cells
        h[0].text = "Fuente"
        h[1].text = "Artículos"
        for k, v in sorted(fuentes.items(), key=lambda x: -x[1]):
            r = t.add_row().cells
            r[0].text = k
            r[1].text = str(v)

    # Pie metodológico
    doc.add_heading("8. Notas metodológicas", level=1)
    doc.add_paragraph(
        "El score global agrega seis indicadores (estabilidad gubernamental, conflictos sociales, "
        "riesgo regulatorio, polarización, corrupción y seguridad) ponderados según pesos definidos "
        "en config.yaml. El sentimiento se calcula con un lexicón en español; en producción se "
        "recomienda reemplazarlo por un modelo transformer (p. ej. pysentimiento/BETO). Las entidades "
        "se extraen mediante un diccionario curado de actores políticos peruanos. El listado de "
        "fuentes RSS, portales del Estado y datasets internacionales (GDELT/ACLED) figura en "
        "config.yaml. En modo 'demo' se usan datos sintéticos representativos."
    )

    doc.save(output_path)
    return output_path


def _generar_hallazgos(temas, conflictos, entidades, riesgo):
    hallazgos = []
    cats = sorted(riesgo["categorias"].items(), key=lambda x: -x[1])
    if cats:
        cat_top, sc_top = cats[0]
        hallazgos.append(
            f"Categoría con mayor riesgo: {cat_top.replace('_',' ')} (score {sc_top}/100). "
            "Recomendación: priorizar monitoreo y briefings semanales sobre esta dimensión."
        )

    activos_alta = [c for c in conflictos if (c.raw or {}).get("severidad") == "alta" and (c.raw or {}).get("estado") == "activo"]
    if activos_alta:
        regiones = sorted({(c.raw or {}).get("region") or c.region or "—" for c in activos_alta})
        hallazgos.append(
            f"{len(activos_alta)} conflictos sociales activos de severidad ALTA en: "
            f"{', '.join(regiones)}. Riesgo de paralización de operaciones extractivas y vías clave."
        )

    insts = entidades.get("instituciones", [])
    if insts:
        top_i = insts[0][0]
        hallazgos.append(
            f"Institución más expuesta en cobertura: {top_i}. Implica mayor escrutinio y foco "
            "para mensajes de comunicación política."
        )

    if riesgo["sentimiento_promedio"] < -0.2:
        hallazgos.append(
            "Tono general de la cobertura es marcadamente negativo. Se recomienda activar plan de "
            "narrativa contraria y priorizar mensajes de estabilidad institucional."
        )

    if not hallazgos:
        hallazgos.append("No se detectan señales críticas en la ventana analizada.")
    return hallazgos
