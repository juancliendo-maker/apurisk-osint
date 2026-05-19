"""Extracción de texto desde documentos adjuntos por el analista.

Soporta:
  - PDF (con pdfplumber)
  - DOCX (con python-docx)
  - TXT, MD (lectura directa)

El texto extraído se inyecta al motor analítico como contexto adicional
para enriquecer el matching de keywords y mejorar la calidad del reporte.

Límites de seguridad:
  - Máx 20 MB por archivo (configurable)
  - Máx 50,000 caracteres extraídos por documento (truncado si excede)
  - Solo se aceptan tipos MIME conocidos / extensiones seguras
"""
from __future__ import annotations
import io
from typing import Optional


MAX_BYTES_PER_FILE = 20 * 1024 * 1024     # 20 MB
MAX_CHARS_PER_DOC = 50_000                  # 50K caracteres
ACCEPTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}


def extract_from_pdf(file_bytes: bytes) -> str:
    """Extrae texto de un PDF usando pdfplumber.

    Si pdfplumber no está disponible, retorna mensaje de error y string vacío.
    """
    try:
        import pdfplumber
    except ImportError:
        return ""

    try:
        text_pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                if i >= 50:  # max 50 páginas para no saturar
                    break
                try:
                    t = page.extract_text() or ""
                    if t.strip():
                        text_pages.append(t)
                except Exception:
                    continue
        return "\n\n".join(text_pages)
    except Exception as e:
        print(f"  [warn] extract_from_pdf: {e}")
        return ""


def extract_from_docx(file_bytes: bytes) -> str:
    """Extrae texto de un .docx usando python-docx."""
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        # Tablas también
        for table in doc.tables:
            for row in table.rows:
                celdas = [c.text.strip() for c in row.cells if c.text.strip()]
                if celdas:
                    parrafos.append(" | ".join(celdas))
        return "\n".join(parrafos)
    except Exception as e:
        print(f"  [warn] extract_from_docx: {e}")
        return ""


def extract_from_text(file_bytes: bytes) -> str:
    """Lee archivos de texto plano (TXT, MD)."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _get_extension(filename: str) -> str:
    if not filename:
        return ""
    idx = filename.rfind(".")
    if idx == -1:
        return ""
    return filename[idx:].lower()


def extract_document(filename: str, content_type: str,
                       file_bytes: bytes) -> dict:
    """Despacha la extracción según extensión y MIME type.

    Returns dict con:
      - nombre: nombre del archivo
      - texto_extraido: texto plano del documento (truncado a MAX_CHARS)
      - caracteres: número de caracteres extraídos
      - tipo: pdf | docx | txt | md | unsupported
      - error: mensaje de error si aplica
    """
    result = {
        "nombre": filename or "documento",
        "texto_extraido": "",
        "caracteres": 0,
        "tipo": "unsupported",
        "error": "",
    }

    if not file_bytes:
        result["error"] = "Archivo vacío"
        return result

    if len(file_bytes) > MAX_BYTES_PER_FILE:
        result["error"] = (f"Archivo excede el límite de "
                            f"{MAX_BYTES_PER_FILE // (1024*1024)} MB")
        return result

    ext = _get_extension(filename)
    ct = (content_type or "").lower()

    texto = ""
    if ext == ".pdf" or "pdf" in ct:
        result["tipo"] = "pdf"
        texto = extract_from_pdf(file_bytes)
    elif ext == ".docx" or "wordprocessingml" in ct or "word" in ct:
        result["tipo"] = "docx"
        texto = extract_from_docx(file_bytes)
    elif ext in (".txt", ".md", ".markdown") or "text/" in ct:
        result["tipo"] = "txt" if ext == ".txt" else "md"
        texto = extract_from_text(file_bytes)
    else:
        result["error"] = (f"Tipo de archivo no soportado: ext={ext} "
                            f"content_type={ct}. Soportados: PDF, DOCX, TXT, MD.")
        return result

    # Truncar si excede max
    if len(texto) > MAX_CHARS_PER_DOC:
        texto = texto[:MAX_CHARS_PER_DOC] + f"\n\n[...truncado a {MAX_CHARS_PER_DOC} caracteres]"

    if not texto.strip():
        result["error"] = "No se pudo extraer texto del documento (¿es escaneado o protegido?)."

    result["texto_extraido"] = texto.strip()
    result["caracteres"] = len(result["texto_extraido"])
    return result
